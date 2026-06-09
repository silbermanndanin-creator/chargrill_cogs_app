"""Part-time variation letters.

(1) detect_variations: compare a week's actual Tanda shifts to each part-timer's
    contracted pattern. A variation is flagged when the actual START time differs from
    the contracted start by more than START_TOL_MIN, or when a shift is worked on a
    non-contracted day. End-time differences are ignored (owner: 30-40 min is normal).
(2) combine_patterns: group an employee's variation events (across one or many weeks)
    by weekday + start time, so a recurring change becomes ONE letter spanning a date
    range instead of many near-identical weekly letters.
(3) render_letter: fill the firm's 'Variation of Employment Agreement' .docx template.
"""
import io
import os
import datetime as dt

import pandas as pd

import contracts as C
from payroll import load_csv_from_bytes, parse_hhmm

START_TOL_MIN = 15  # start-time differences <= this many minutes are ignored
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "variation_letter.docx")


# ── time helpers ──────────────────────────────────────────────────────────────
def _mins(hhmm):
    t = parse_hhmm(hhmm)
    return t.hour * 60 + t.minute if t else None


def _fmt(hhmm):
    """'14:00' -> '2:00pm'."""
    t = parse_hhmm(hhmm)
    if not t:
        return str(hhmm)
    ap = "am" if t.hour < 12 else "pm"
    return f"{t.hour % 12 or 12}:{t.minute:02d}{ap}"


def _fmt_compact(hhmm):
    """'14:00' -> '2pm', '14:30' -> '2:30pm' (drop ':00' on the hour)."""
    t = parse_hhmm(hhmm)
    if not t:
        return str(hhmm)
    ap = "am" if t.hour < 12 else "pm"
    h = t.hour % 12 or 12
    return f"{h}{ap}" if t.minute == 0 else f"{h}:{t.minute:02d}{ap}"


def _ordinal(n):
    """1 -> '1st', 2 -> '2nd', 11 -> '11th'."""
    suf = "th" if 11 <= n % 100 <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suf}"


def nice_date(d):
    """A date -> 'Monday 1st June'."""
    if not isinstance(d, dt.date):
        try:
            d = pd.to_datetime(d).date()
        except Exception:
            return str(d)
    return f"{d:%A} {_ordinal(d.day)} {d:%B}"


# Letters round a finish in this window up to 10:00pm (a 9:35pm finish is really the 10pm
# shift). Display keeps the actual time; only the letter is rounded.
LETTER_FINISH_WINDOW = (21 * 60 + 25, 22 * 60)  # 9:25pm … 10:00pm


def _letter_finish(finish):
    m = _mins(finish)
    lo, hi = LETTER_FINISH_WINDOW
    return "22:00" if (m is not None and lo <= m <= hi) else finish


def _merge_day(events):
    """Merge ONE day's CONTIGUOUS/overlapping blocks into spans — a split that's really one
    continuous shift (11am-3pm + 3pm-9:35pm -> 11am-9:35pm) becomes one. Blocks with a real
    gap between them are kept separate. Returns events with the same keys."""
    evs = sorted([e for e in events if _mins(e["actual_start"]) is not None],
                 key=lambda e: _mins(e["actual_start"]))
    out = []
    for e in evs:
        s, f = _mins(e["actual_start"]), _mins(e["actual_finish"])
        if out and f is not None and out[-1]["_f"] is not None and s <= out[-1]["_f"]:
            prev = out[-1]  # touches/overlaps the previous block -> extend the span
            if f > prev["_f"]:
                prev["_f"], prev["actual_finish"] = f, e["actual_finish"]
            if e["kind"] == "start" and prev["kind"] != "start":
                prev["kind"], prev["contracted_start"] = "start", e.get("contracted_start")
        else:
            out.append({"date": e["date"], "weekday": e["weekday"],
                        "actual_start": e["actual_start"], "actual_finish": e["actual_finish"],
                        "contracted_start": e.get("contracted_start"),
                        "contracted_finish": e.get("contracted_finish"),
                        "kind": e["kind"], "_f": f if f is not None else s})
    for o in out:
        o.pop("_f", None)
    return out


def merge_events(events):
    """Merge each day's contiguous shift blocks across all of an employee's events (any number
    of days/weeks). Feeds combine_patterns so letters reflect one span per continuous shift."""
    by_date = {}
    for e in events:
        d = e["date"] if isinstance(e["date"], dt.date) else pd.to_datetime(e["date"]).date()
        by_date.setdefault(d, []).append({**e, "date": d})
    out = []
    for d in sorted(by_date):
        out.extend(_merge_day(by_date[d]))
    return out


def _day_blocks(events):
    """A day's worked time for the table: contiguous blocks merged ('11am-9:35pm'), genuinely
    separate shifts shown apart ('11am-3pm, 5pm-9pm'). Actual times (no letter rounding)."""
    parts = []
    for m in _merge_day(events):
        s = _fmt_compact(m["actual_start"])
        parts.append(f"{s}–{_fmt_compact(m['actual_finish'])}"
                     if _mins(m["actual_finish"]) is not None else s)
    return ", ".join(parts) if parts else "—"


def display_rows(vmap):
    """Merge a week's variation events into ONE row per employee+date for the on-screen table:
    {Employee, When ('Monday 1st June'), Worked ('11am-3pm'), 'Contracted start', Type}.
    A split shift (two entries on one day) becomes a single 'earliest-start to latest-finish' span."""
    rows = []
    for emp, evs in vmap.items():
        by_date = {}
        for e in evs:
            d = e["date"] if isinstance(e["date"], dt.date) else pd.to_datetime(e["date"]).date()
            by_date.setdefault(d, []).append(e)
        for d in sorted(by_date):
            grp = by_date[d]
            cstart = next((g.get("contracted_start") for g in grp if g.get("contracted_start")), None)
            rows.append({
                "Employee": emp,
                "When": nice_date(d),
                "Worked": _day_blocks(grp),
                "Contracted start": _fmt_compact(cstart) if cstart else "—",
                "Type": "start time" if any(g["kind"] == "start" for g in grp) else "extra day",
            })
    return rows


def _dur(start, finish):
    s, f = _mins(start), _mins(finish)
    if s is None or f is None:
        return 0.0
    if f < s:
        f += 24 * 60
    return (f - s) / 60.0


def _join(items):
    items = list(items)
    if len(items) <= 1:
        return "".join(items)
    return ", ".join(items[:-1]) + " and " + items[-1]


# ── detection ─────────────────────────────────────────────────────────────────
def detect_variations(shift_df, cmap, start_tol_min=START_TOL_MIN):
    """{canonical_name: [event, ...]} for tracked part-timers. cmap is the contracts
    map {name: {weekday: (start, finish)}} from storage.load_contracts(). Each event:
    {date, weekday, contracted_start, actual_start, actual_finish, contracted_finish, kind}.
    kind = 'start' (started materially off-contract) or 'extra_day' (non-contracted day)."""
    out = {}
    if shift_df is None or shift_df.empty or not cmap:
        return out
    key_col = "_name_key" if "_name_key" in shift_df else "Name"
    for _, g in shift_df.groupby(key_col):
        raw_name = str(g["Name"].iloc[0])
        cname, days = C.match_contract(raw_name, cmap)
        if not days:
            continue  # not a tracked part-timer
        events = []
        for _, row in g.iterrows():
            d = pd.to_datetime(row["Date"]).date()
            wd = C.IDX_TO_DAY[d.weekday()]
            astart = str(row["Shift Start Time"]).strip()
            afin = str(row["Shift End Time"]).strip()
            if wd in days:
                cstart, cfin = days[wd]
                am, cm = _mins(astart), _mins(cstart)
                if am is not None and cm is not None and abs(am - cm) > start_tol_min:
                    events.append({"date": d, "weekday": wd, "contracted_start": cstart,
                                   "actual_start": astart, "actual_finish": afin,
                                   "contracted_finish": cfin, "kind": "start"})
            else:
                events.append({"date": d, "weekday": wd, "contracted_start": None,
                               "actual_start": astart, "actual_finish": afin,
                               "contracted_finish": None, "kind": "extra_day"})
        if events:
            out[cname] = sorted(events, key=lambda e: e["date"])
    return out


# ── combine recurring variations across weeks ──────────────────────────────────
def combine_patterns(events):
    """Group events (possibly spanning several weeks) by (weekday, actual_start,
    actual_finish) into patterns. Each pattern: {weekday, start, finish, kind, dates[]}."""
    groups = {}
    for e in events:
        k = (e["weekday"], e["actual_start"], e["actual_finish"])
        g = groups.setdefault(k, {"weekday": e["weekday"], "start": e["actual_start"],
                                  "finish": e["actual_finish"], "kind": e["kind"],
                                  "contracted_start": e.get("contracted_start"),
                                  "contracted_finish": e.get("contracted_finish"),
                                  "dates": []})
        d = e["date"] if isinstance(e["date"], dt.date) else pd.to_datetime(e["date"]).date()
        g["dates"].append(d)
    patterns = []
    for g in groups.values():
        g["dates"] = sorted(g["dates"])
        patterns.append(g)
    # order by first occurrence then weekday
    patterns.sort(key=lambda p: (p["dates"][0], C.WEEKDAYS[p["weekday"]]))
    return patterns


# ── docx letter rendering ───────────────────────────────────────────────────────
def _set(p, text):
    """Replace a paragraph's text in-place, keeping its paragraph style."""
    runs = list(p.runs)
    for r in runs[1:]:
        r.text = ""
    if runs:
        runs[0].text = text
    else:
        p.add_run(text)


def _del(p):
    p._element.getparent().remove(p._element)


def _name_sub(text, cname):
    for tok in ("[Insert name of employee]", "[insert name of employee]"):
        text = text.replace(tok, cname)
    return text


def render_letter(cname, patterns, today=None, details=None):
    """Fill the template for one employee + a set of (possibly combined) patterns.
    Returns .docx bytes. Fields we can't know (address, agreement date, signatory,
    meal breaks, return-to/by) stay as [insert ...] placeholders unless provided in
    contracts.EMP_DETAILS — the letter is a DRAFT for the owner to review and sign."""
    from docx import Document
    today = today or dt.date.today()
    details = details or {}
    doc = Document(TEMPLATE_PATH)

    commence = min(min(p["dates"]) for p in patterns)
    end = max(max(p["dates"]) for p in patterns)
    days_full, seen = [], set()
    for p in patterns:
        f = C.WEEKDAY_FULL[p["weekday"]]
        if f not in seen:
            seen.add(f)
            days_full.append(f)
    # Clause (iii): the template's exact sentence, repeated per pattern for split shifts.
    # Finish times are rounded up to 10pm in the 9:25–10pm window (see _letter_finish).
    iii = "; ".join(f"you will start work at {_fmt_compact(p['start'])} and finish work at "
                    f"{_fmt_compact(_letter_finish(p['finish']))} on {C.WEEKDAY_FULL[p['weekday']]}"
                    for p in patterns)
    # Clause (i): ordinary hours per day = sum of that day's blocks. Fill only if every
    # worked day has the same daily total; otherwise leave the [insert] blank for the owner.
    day_hours = {}
    for p in patterns:
        day_hours[p["weekday"]] = day_hours.get(p["weekday"], 0.0) + _dur(p["start"], _letter_finish(p["finish"]))
    distinct = {round(v, 2) for v in day_hours.values()}
    hours_str = f"{next(iter(distinct)):g}" if len(distinct) == 1 else "[insert]"

    to_delete = []
    for p in doc.paragraphs:
        raw = p.text
        t = raw.strip()
        if not t:
            continue
        if t == "[Insert date]":
            _set(p, today.strftime("%d %B %Y")); continue
        if t.startswith("[Option 1"):
            to_delete.append(p); continue
        if t.startswith("[Option 2"):  # template wording verbatim, dates filled, instruction prefix dropped
            _set(p, f"The terms of this variation will commence on {commence:%d %B %Y}, and will "
                    f"end on {end:%d %B %Y}. At the end of this period, your hours of work and "
                    f"associated arrangements will revert back to those set out in [clause 3] of the "
                    f"Agreement (instead of those set out in clause [3] below)."); continue
        if "ordinary hours each day" in raw:
            _set(p, f"(i)\tyou will work {hours_str} ordinary hours each day;"); continue
        if "days of the week to be worked" in raw:
            _set(p, f"(ii)\tyou will work on {_join(days_full)} each week;"); continue
        if "you will start work at" in raw:
            _set(p, f"(iii)\t{iii};"); continue
        if "address details of employee" in raw and details.get("address"):
            _set(p, details["address"]); continue
        new = raw
        if details.get("agreement_date"):
            new = new.replace("[insert date of Employment Agreement]", details["agreement_date"])
        if details.get("return_to"):
            new = new.replace("[insert name]", details["return_to"])
        if details.get("return_by") and "return it to" in raw:
            new = new.replace("[insert date]", details["return_by"])
        if details.get("signatory"):
            new = new.replace("[Insert signatory]", details["signatory"])
        new = _name_sub(new, cname)
        if new != raw:
            _set(p, new)
    for p in to_delete:
        _del(p)

    # name token inside the signature table
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "name of employee" in p.text:
                        _set(p, _name_sub(p.text, cname))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def summarise(patterns):
    """Short human summary of an employee's variation patterns, for the history list."""
    bits = []
    for p in patterns:
        wd = C.WEEKDAY_FULL[p["weekday"]]
        n = len(p["dates"])
        span = (f"{p['dates'][0]:%d %b}" if n == 1
                else f"{p['dates'][0]:%d %b}–{p['dates'][-1]:%d %b} ×{n}")
        if p["kind"] == "extra_day":
            bits.append(f"{wd} extra ({_fmt(p['start'])}) {span}")
        else:
            bits.append(f"{wd} {_fmt(p['contracted_start'])}→{_fmt(p['start'])} {span}")
    return " · ".join(bits)
